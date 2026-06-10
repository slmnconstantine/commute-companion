import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

def correct_docx(filepath, output_filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    doc = Document(filepath)
    
    # List of replacements per paragraph
    # Format: paragraph_index, [ (old_text, new_text), ... ]
    replacements = {
        32: [
            ("Seng et al. (2023)", "Seng, Ang, Ngharamike, and Peter (2023)")
        ],
        33: [
            ("Li et al. (2022)", "Li, Gao, Wang, Huang, and Nie (2022)")
        ],
        35: [
            ("Hashikami et al. (2023)", "Hashikami, Kobayashi, Li, Nakano, and Shigeno (2023)")
        ],
        36: [
            ("Rao et al. (2025)", "Rao, Dalal, Agarwal, Calacci, and Monroy-Hernández (2025)")
        ],
        42: [
            ("Dastani et al. (2024)", "Dastani, Koosha, Karimi, and Moghaddam (2024)"),
            ("Psaraftis et al. (2024)", "Psaraftis, Ntalianis, and Mastorakis (2024)"),
            ("Shah et al. (2020)", "Shah, El Affendi, and Qureshi (2020)"),
            ("Peng et al. (2024)", "Peng, Wang, Wang, and Yu (2026)")
        ],
        43: [
            ("Afsari et al. (2025)", "Afsari, Salehi, Miristice, and Gentile (2026)")
        ],
        44: [
            ("Zafar et al. (2022)", "Zafar, Khattak, Aloqaily, and Hussain (2022)"),
            ("Mitropoulos et al. (2021)", "Mitropoulos, Kortsari, and Ayfantopoulou (2021)")
        ]
    }
    
    for idx, reps in replacements.items():
        p = doc.paragraphs[idx]
        print(f"Updating paragraph {idx}...")
        for old, new in reps:
            if old in p.text:
                print(f"  Replacing: '{old}' -> '{new}'")
                # Attempt to replace inside runs first
                replaced = False
                for r in p.runs:
                    if old in r.text:
                        r.text = r.text.replace(old, new)
                        replaced = True
                        break
                
                # Fallback: if split across runs, merge them
                if not replaced:
                    print(f"  Note: Citation split across runs. Merging paragraph runs.")
                    full_text = p.text.replace(old, new)
                    # Clear runs and write to first run
                    for r in p.runs:
                        r.text = ""
                    if p.runs:
                        p.runs[0].text = full_text
                    else:
                        p.add_run(full_text)
            else:
                print(f"  WARNING: '{old}' not found in paragraph {idx}!")
                
    doc.save(output_filepath)
    print(f"Successfully saved corrected manuscript to {output_filepath}")

if __name__ == "__main__":
    correct_docx("COMMUTE COMPANION Manuscript_Formatted.docx", "COMMUTE COMPANION Manuscript_Formatted.docx")
