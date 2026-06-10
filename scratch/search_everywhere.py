import os
from docx import Document

def search_everywhere(filepath, target):
    if not os.path.exists(filepath):
        print(f"Not found: {filepath}")
        return
    doc = Document(filepath)
    print(f"=== Searching in {filepath} for '{target}' ===")
    
    # Search in paragraphs
    for i, p in enumerate(doc.paragraphs):
        if target.lower() in p.text.lower():
            print(f"Paragraph {i}: {p.text}")
            
    # Search in tables
    for i, table in enumerate(doc.tables):
        for r_idx, row in enumerate(table.rows):
            for c_idx, cell in enumerate(row.cells):
                for p_idx, p in enumerate(cell.paragraphs):
                    if target.lower() in p.text.lower():
                        print(f"Table {i}, Row {r_idx}, Cell {c_idx}, Paragraph {p_idx}: {p.text}")

if __name__ == "__main__":
    search_everywhere("COMMUTE COMPANION Manuscript.docx", "Device Status")
    search_everywhere("COMMUTE COMPANION Manuscript.docx", "Home Screen")
    search_everywhere("COMMUTE COMPANION Manuscript.docx", "Figure 6")
