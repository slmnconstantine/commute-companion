import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

def print_early(filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    doc = Document(filepath)
    search_terms = [
        "seng", "li", "golightly", "hashikami", "rao", "anagnostopoulos",
        "dastani", "psaraftis", "shah", "peng", "afsari", "gkartzonikas",
        "makhdomi", "zafar", "mitropoulos"
    ]
    
    for idx in range(0, min(67, len(doc.paragraphs))):
        p = doc.paragraphs[idx]
        text_lower = p.text.lower()
        found = [term for term in search_terms if term in text_lower]
        if found:
            print(f"[{idx}]: {p.text}\n")

if __name__ == "__main__":
    print_early("COMMUTE COMPANION Manuscript_Formatted.docx")
