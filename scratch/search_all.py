import os
from docx import Document

def search_files(target):
    root_dir = r"c:\Users\iansa\Desktop\System\commute-companion"
    for root, dirs, files in os.walk(root_dir):
        if ".git" in root or "node_modules" in root or ".expo" in root:
            continue
        for file in files:
            # Skip temp office files
            if file.startswith("~$"):
                continue
            filepath = os.path.join(root, file)
            if file.endswith(".docx"):
                try:
                    doc = Document(filepath)
                    for i, p in enumerate(doc.paragraphs):
                        if target in p.text.lower():
                            print(f"[Docx Paragraph] {filepath} | Index {i}: {p.text}")
                    for t_idx, table in enumerate(doc.tables):
                        for r_idx, row in enumerate(table.rows):
                            for c_idx, cell in enumerate(row.cells):
                                for p_idx, p in enumerate(cell.paragraphs):
                                    if target in p.text.lower():
                                        print(f"[Docx Table] {filepath} | Table {t_idx}, Row {r_idx}, Cell {c_idx}: {p.text}")
                except Exception as e:
                    print(f"Error reading docx {filepath}: {e}")
            elif file.endswith((".txt", ".md", ".json", ".js", ".ts", ".html", ".css")):
                try:
                    with open(filepath, "r", encoding="utf-8", errors="ignore") as f:
                        content = f.read()
                        if target in content.lower():
                            print(f"[Text File] {filepath}")
                except Exception as e:
                    print(f"Error reading text {filepath}: {e}")

if __name__ == "__main__":
    search_files("weather status")
    search_files("datm")
