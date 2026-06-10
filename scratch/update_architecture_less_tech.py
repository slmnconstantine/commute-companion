import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

def update_architecture_description(filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    doc = Document(filepath)
    
    new_description = (
        "The system architecture of Commute Companion is composed of four interconnected layers "
        "that coordinate to deliver a seamless ride-sharing experience. At the front end is the mobile "
        "application that users interact with directly, which handles the user interface, manages the "
        "app's local settings, and captures user actions like GPS tracking and voice commands. This "
        "application communicates with the rest of the system through a communication layer, which acts "
        "as a bridge for transferring standard data requests (such as logging in or uploading verification "
        "documents) and establishing real-time connections to stream live driver locations to riders. The "
        "core of the system is powered by backend services, which securely store user accounts, manage "
        "uploaded identification documents, and execute server-side operations such as handling voice "
        "commands. Finally, the system integrates several external services to handle specialized tasks, "
        "including an artificial intelligence model to process voice commands, mapping services for navigation "
        "and route planning, notification gateways to send real-time alerts to user devices, and a payment "
        "processing gateway to handle transactions."
    )
    
    found = False
    for idx, p in enumerate(doc.paragraphs):
        # Match paragraph containing the previous technical description or the placeholder
        if "operational blocks" in p.text and "Supabase" in p.text:
            print(f"Found technical description at paragraph index {idx}.")
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = new_description
            else:
                p.add_run(new_description)
            found = True
            break
            
    if found:
        doc.save(filepath)
        print("Successfully updated the manuscript with the less technical description.")
    else:
        print("WARNING: Could not find the target paragraph to update!")

if __name__ == "__main__":
    update_architecture_description("COMMUTE COMPANION Manuscript_Formatted.docx")
