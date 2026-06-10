import os
import sys
from docx import Document

# Configure stdout to use utf-8
sys.stdout.reconfigure(encoding='utf-8')

def print_references(filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    doc = Document(filepath)
    print(f"Total paragraphs: {len(doc.paragraphs)}")
    
    # Let's find paragraphs containing "references" or "bibliography"
    ref_idx = -1
    for i, p in enumerate(doc.paragraphs):
        text = p.text.strip().lower()
        if text == "references" or text == "bibliography":
            print(f"Found references section at paragraph index {i}: {p.text}")
            ref_idx = i
            break
            
    if ref_idx != -1:
        # Print next 100 paragraphs
        print("=== References ===")
        for j in range(ref_idx + 1, min(ref_idx + 100, len(doc.paragraphs))):
            print(f"[{j}]: {doc.paragraphs[j].text}")
    else:
        # Try searching for some of the authors in paragraphs
        print("Could not find exact References/Bibliography heading. Searching for 'Seng' or 'Golightly' or 'Hashikami'...")
        for i, p in enumerate(doc.paragraphs):
            if any(name in p.text for name in ["Seng", "Golightly", "Hashikami", "Rao", "Afsari", "Gkartzonikas"]):
                print(f"[{i}]: {p.text}")

if __name__ == "__main__":
    print_references("COMMUTE COMPANION Manuscript_Formatted.docx")
