import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

def print_range(filepath, start, end):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    doc = Document(filepath)
    for idx in range(start, min(end, len(doc.paragraphs))):
        print(f"[{idx}]: {doc.paragraphs[idx].text}\n")

if __name__ == "__main__":
    print_range("COMMUTE COMPANION Manuscript_Formatted.docx", 82, 93)
