import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

def find_citations(filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    doc = Document(filepath)
    search_terms = [
        "seng", "li", "golightly", "hashikami", "rao", "anagnostopoulos",
        "dastani", "psaraftis", "shah", "peng", "afsari", "gkartzonikas",
        "makhdomi", "zafar", "mitropoulos"
    ]
    
    print(f"Searching early paragraphs in {filepath}...")
    for idx in range(0, min(100, len(doc.paragraphs))):
        p = doc.paragraphs[idx]
        text_lower = p.text.lower()
        found = [term for term in search_terms if term in text_lower]
        if found:
            print(f"Paragraph {idx} (Found: {found}):")
            print(f"  {p.text}\n")

if __name__ == "__main__":
    find_citations("COMMUTE COMPANION Manuscript_Formatted.docx")
