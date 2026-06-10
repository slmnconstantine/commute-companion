import os
import sys
from docx import Document

sys.stdout.reconfigure(encoding='utf-8')

def update_architecture_description(filepath):
    if not os.path.exists(filepath):
        print(f"File not found: {filepath}")
        return
    doc = Document(filepath)
    
    target_placeholder = "This diagram illustrates the system structure, including the app, the backend, and the external services."
    
    new_description = (
        "The system architecture of the Commute Companion platform is structured into four main operational blocks: "
        "the Client Application, the Communication Layer, the Supabase Backend Services, and External Services. "
        "The Client Application, developed using React Native and Expo (v56), is operated by the user and consists of Frontend UI Screens "
        "and Core State & Contexts (managing the local database, routing, voice assistant, push notifications, and live tracking). "
        "The Communication Layer manages data transfer, utilizing HTTPS REST API Channels for authentication, database updates, "
        "document uploads, and voice file transfers, alongside WebSockets for streaming real-time driver GPS coordinates. "
        "The Supabase Backend Services handle central processes, including database storage and authentication, file storage for user "
        "avatars and verification documents, and Edge Functions that process voice-activated commands. "
        "Finally, External Services provide specialized capabilities, including the Grok API (xAI) for audio translation and natural language processing (NLP) intent parsing, "
        "the MapLibre and OSRM APIs for navigation and routing, the Expo Push Service and Firebase Cloud Messaging (FCM) for push notifications, "
        "and the PayMongo API for payment processing."
    )
    
    found = False
    for idx, p in enumerate(doc.paragraphs):
        if p.text.strip() == target_placeholder:
            print(f"Found placeholder at paragraph index {idx}.")
            # Clear all runs except the first, or write to first run
            for r in p.runs:
                r.text = ""
            if p.runs:
                p.runs[0].text = new_description
            else:
                p.add_run(new_description)
            found = True
            break
            
    if found:
        doc.save(filepath)
        print("Successfully updated the manuscript with the new description.")
    else:
        print("WARNING: Could not find the target placeholder paragraph in the document!")

if __name__ == "__main__":
    update_architecture_description("COMMUTE COMPANION Manuscript_Formatted.docx")
