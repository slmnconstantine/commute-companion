import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.section import WD_SECTION

def format_document(input_path, output_path):
    doc = Document(input_path)

    # 1.4. Paper Size: 8.5" x 11"
    # 1.3. Margins: Left 1.5 in, Top/Right/Bottom 1.0 in
    # 1.6 & 1.7: Pagination in Top Right corner, no page number on 1st page
    for section in doc.sections:
        section.page_width = Inches(8.5)
        section.page_height = Inches(11.0)
        section.left_margin = Inches(1.5)
        section.right_margin = Inches(1.0)
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        
        # Pagination: different first page header
        section.different_first_page_header_footer = True
        
        # Add page numbers to header
        # Note: python-docx doesn't natively support adding dynamic page numbers easily 
        # without raw XML. We will inject the page number XML.
        header = section.header
        
        # Clear existing paragraphs in header
        for p in header.paragraphs:
            p.text = ""
            
        header_para = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
        header_para.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        
        # Adding a page number field requires raw XML
        from docx.oxml.shared import OxmlElement, qn
        
        run = header_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(qn('w:fldCharType'), 'begin')
        instrText = OxmlElement('w:instrText')
        instrText.set(qn('xml:space'), 'preserve')
        instrText.text = "PAGE"
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(qn('w:fldCharType'), 'separate')
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(qn('w:fldCharType'), 'end')
        
        run._r.append(fldChar1)
        run._r.append(instrText)
        run._r.append(fldChar2)
        run._r.append(fldChar3)
        
        # Set font for header run
        run.font.name = 'Times New Roman'
        run.font.size = Pt(12)
        
        # Make sure first page header is empty
        first_page_header = section.first_page_header
        for p in first_page_header.paragraphs:
            p.text = ""

    # Change style defaults
    # 1.1 Font: Times New Roman, 12pt
    # 1.2 Spacing: 1.5-line space
    # 1.3 Paragraphs: Justified and indented
    for paragraph in doc.paragraphs:
        # Paragraph alignment justified
        # We only justify normal paragraphs or body text. Let's do it for all that aren't centered/right
        if paragraph.alignment is None or paragraph.alignment == WD_ALIGN_PARAGRAPH.LEFT:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        
        # Indent first line (approx 0.5 inch is standard)
        paragraph.paragraph_format.first_line_indent = Inches(0.5)
        
        # 1.5 line spacing
        paragraph.paragraph_format.line_spacing = 1.5
        
        for run in paragraph.runs:
            run.font.name = 'Times New Roman'
            run.font.size = Pt(12)

    # Note: tables might also need font updates
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for run in paragraph.runs:
                        run.font.name = 'Times New Roman'
                        run.font.size = Pt(12)

    doc.save(output_path)

if __name__ == "__main__":
    format_document("COMMUTE COMPANION Manuscript.docx", "COMMUTE COMPANION Manuscript_Formatted.docx")
    print("Formatting complete.")
